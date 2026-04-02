import pytest
import sys
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from PyQt6.QtWidgets import QApplication, QLineEdit, QComboBox, QTextEdit
from docxtpl import DocxTemplate
from docx import Document


# Ensure QApplication exists for widget tests
@pytest.fixture(scope="module")
def qapp():
    app = QApplication.instance() or QApplication(sys.argv)
    yield app


LINE_EDIT_FIELDS = [
    "project_title", "project_address", "nya_project_code", "client_project_code",
    "client", "client_name", "client_title", "client_license",
    "client_phone", "client_mobile", "client_email",
    "client_office_no", "client_invoice_email",
    "request_date", "work_type", "fee", "save_location", "a250_creator",
]
COMBO_FIELDS = ["principal_name", "project_manager", "fee_type"]
MULTILINE_FIELDS = ["project_description", "detailed_scope", "client_address", "invoice_to"]


class TestA250Generation:

    def _make_a250_vars(self, overrides=None):
        """Build a minimal a250_vars dict matching current widget types."""
        vars_ = {}
        for f in LINE_EDIT_FIELDS:
            m = Mock(spec=QLineEdit)
            m.text = Mock(return_value="")
            vars_[f] = m
        for f in COMBO_FIELDS:
            m = Mock(spec=QComboBox)
            m.currentText = Mock(return_value="")
            vars_[f] = m
        for f in MULTILINE_FIELDS:
            m = Mock(spec=QTextEdit)
            m.toPlainText = Mock(return_value="")
            vars_[f] = m
        vars_["project_title"].text = Mock(return_value="MyProject")
        if overrides:
            for k, v in overrides.items():
                if k in vars_:
                    widget = vars_[k]
                    if hasattr(widget, 'toPlainText'):
                        widget.toPlainText = Mock(return_value=v)
                    elif hasattr(widget, 'currentText'):
                        widget.currentText = Mock(return_value=v)
                    else:
                        widget.text = Mock(return_value=v)
        return vars_

    def test_render_called_with_all_fields(self, qapp, tmp_path):
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({"save_location": str(tmp_path)})
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("app.QMessageBox.critical"):
                with patch("subprocess.Popen"):
                    window._generate_a250(a250_vars)
        render_call = mock_doc.render.call_args[0][0]
        assert "project_title" in render_call
        assert "current_date" in render_call
        assert "requested_by" in render_call
        assert "client_signed" in render_call
        assert "invoice_to" in render_call
        assert render_call["project_title"] == "MyProject"

    def test_output_filename_uses_project_title(self, qapp, tmp_path):
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({"save_location": str(tmp_path)})
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        save_call = mock_doc.save.call_args[0][0]
        assert "MyProject" in str(save_call)
        assert str(save_call).endswith(".docx")

    def test_save_location_used_when_provided(self, qapp, tmp_path):
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({"save_location": str(tmp_path)})
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        save_path = mock_doc.save.call_args[0][0]
        assert str(tmp_path) in str(save_path)

    def test_save_location_blank_falls_back_to_cwd(self, qapp):
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({"save_location": ""})
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        save_path = Path(mock_doc.save.call_args[0][0])
        assert save_path.parent == Path.cwd()

    def test_missing_template_shows_error_dialog(self, qapp):
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars()
        with patch("app.DocxTemplate", side_effect=FileNotFoundError("template not found")):
            with patch("app.QMessageBox.critical") as mock_dialog:
                window._generate_a250(a250_vars)
        mock_dialog.assert_called_once()
        call_args = mock_dialog.call_args[0]
        assert "Error" in call_args[1]

    def test_requested_by_composite_short(self, qapp, tmp_path):
        """Short name+license+title: single \\n before title."""
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({
            "client_name": "Jane Doe",
            "client_license": "PE",
            "client_title": "Director",
            "client": "Acme Corp",
            "save_location": str(tmp_path),
        })
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        render_data = mock_doc.render.call_args[0][0]
        req_by = render_data["requested_by"]
        assert "Jane Doe" in req_by
        assert "PE" in req_by
        assert "Director" in req_by
        assert "Acme Corp" in req_by
        # Short content: single newline before title (not double)
        assert "\n\n" not in req_by

    def test_requested_by_composite_long(self, qapp, tmp_path):
        """Long name+license+title: double \\n before title."""
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({
            "client_name": "Jonathan Alexander Smithsonian",
            "client_license": "PE, SE, LEED AP BD+C",
            "client_title": "Senior Vice President of Engineering",
            "client": "Big Corporation LLC",
            "save_location": str(tmp_path),
        })
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        render_data = mock_doc.render.call_args[0][0]
        assert "\n\n" in render_data["requested_by"]

    def test_invoice_to_defaults_to_requested_by(self, qapp, tmp_path):
        """Empty invoice_to defaults to requested_by composite."""
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({
            "client_name": "Jane Doe",
            "client_license": "PE",
            "client_title": "Director",
            "client": "Acme Corp",
            "save_location": str(tmp_path),
            # invoice_to left empty (default)
        })
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        render_data = mock_doc.render.call_args[0][0]
        assert render_data["invoice_to"] == render_data["requested_by"]

    def test_invoice_to_custom_text(self, qapp, tmp_path):
        """Non-empty invoice_to uses custom text, not requested_by."""
        from app import FolderSetupApp
        window = FolderSetupApp()
        a250_vars = self._make_a250_vars({
            "invoice_to": "Custom Billing Address\nSuite 100",
            "save_location": str(tmp_path),
        })
        mock_doc = MagicMock()
        with patch("app.DocxTemplate", return_value=mock_doc):
            with patch("subprocess.Popen"):
                window._generate_a250(a250_vars)
        render_data = mock_doc.render.call_args[0][0]
        assert render_data["invoice_to"] == "Custom Billing Address\nSuite 100"


# ---------------------------------------------------------------------------
# Regression tests for A250 template validity (260402-jpl)
# These tests use the real template file (not mocked) to catch Jinja2 syntax
# errors such as the single-brace bug fixed in this quick task.
# ---------------------------------------------------------------------------

TEMPLATE_PATH = Path("templates/A250.docx")

EXPECTED_VARS = {
    "nya_project_code", "project_title", "requested_by", "invoice_to",
    "fee", "today", "client_signed", "principal_name", "project_manager",
}

FULL_RENDER_DATA = {
    "project_title": "Test Project",
    "project_address": "123 Main St",
    "client": "ACME Corp",
    "nya_project_code": "NYA-2026-001",
    "client_project_code": "CP-001",
    "client_name": "Jane Doe",
    "client_title": "PE",
    "client_license": "CA 12345",
    "client_phone": "555-1234",
    "client_mobile": "555-5678",
    "client_email": "jane@acme.com",
    "invoice_to": "Jane Doe\nCA 12345\nPE\nACME Corp",
    "client_office_no": "555-9999",
    "client_invoice_email": "billing@acme.com",
    "client_address": "456 Elm St",
    "request_date": "2026-04-02",
    "work_type": "Structural Review",
    "project_description": "Desc",
    "detailed_scope": "Scope",
    "fee_type": "Lump Sum",
    "fee": "5000",
    "principal_name": "John Smith",
    "project_manager": "Bob Jones",
    "save_location": "",
    "a250_creator": "Tester",
    "today": "April 2, 2026",
    "current_date": "April 2, 2026",
    "requested_by": "Jane Doe\nCA 12345\nPE\nACME Corp",
    "client_signed": "Jane Doe\nPE",
    "received_date": "",
}


class TestA250TemplateRegression:
    """Regression tests ensuring templates/A250.docx remains a valid Jinja2/docxtpl template."""

    def test_template_parses_cleanly(self):
        """Template must parse without TemplateSyntaxError and expose nya_project_code."""
        t = DocxTemplate(str(TEMPLATE_PATH))
        # get_undeclared_template_variables raises TemplateSyntaxError if any tag is malformed
        vars_ = t.get_undeclared_template_variables()
        assert "nya_project_code" in vars_, (
            f"nya_project_code not found in template variables: {sorted(vars_)}"
        )

    def test_expected_vars_present_in_template(self):
        """All core variable names used by _generate_a250 must appear in the template."""
        t = DocxTemplate(str(TEMPLATE_PATH))
        vars_ = t.get_undeclared_template_variables()
        missing = EXPECTED_VARS - vars_
        assert not missing, f"Missing variables in template: {missing}"

    def test_render_produces_file(self, tmp_path):
        """render() + save() with a full data dict must complete without exception."""
        t = DocxTemplate(str(TEMPLATE_PATH))
        t.render(FULL_RENDER_DATA)
        out = tmp_path / "out.docx"
        t.save(str(out))
        assert out.exists(), "Output file was not created"
        assert out.stat().st_size > 0, "Output file is empty"

    def test_render_contains_nya_code(self, tmp_path):
        """The rendered document must contain the nya_project_code value somewhere."""
        t = DocxTemplate(str(TEMPLATE_PATH))
        t.render(FULL_RENDER_DATA)
        out = tmp_path / "out.docx"
        t.save(str(out))
        doc = Document(str(out))
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)
        combined = "\n".join(all_text)
        assert "NYA-2026-001" in combined, (
            "nya_project_code value 'NYA-2026-001' not found in rendered document tables"
        )
